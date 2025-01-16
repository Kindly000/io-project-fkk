import json
import os
from docx import Document
from docx.shared import Inches
from app_backend.logging_f import log_operations_on_file


def format_timestamp(seconds: int) -> str:
    """
    Converts a time duration in seconds into a formatted timestamp string.

    This function takes an integer value representing a time duration in seconds and converts it
    into a human-readable timestamp in the format `[HH:MM:SS]`.

    Args:
        seconds (int): The time duration in seconds to format.

    Returns:
        str: A formatted string representing the time duration in `[HH:MM:SS]` format,
             where `HH` is hours, `MM` is minutes, and `SS` is seconds, all zero-padded to two digits.

    Notes:
        - Handles durations longer than 24 hours (e.g., 90000 seconds -> '[25:00:00]').
        - The input should be a non-negative integer. Negative values may result in unexpected output.

    Example:
        >>> format_timestamp(3661)
        '[01:01:01]'

        >>> format_timestamp(45)
        '[00:00:45]'
    """
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    seconds = seconds % 60

    return f"[{hours:02}:{minutes:02}:{seconds:02}]"


def create_docx_file(note_title: str, note_summary: str,  note_content: list, docx_file_path: str, language: str ='pl') -> bool:
    """
    Creates a DOCX file with the specified note details.

    This function generates a DOCX file containing a note's title, summary, and detailed content.
    The content can include text, speaker labels, and images, formatted with timestamps. The
    language of headings (e.g., "Title" and "Summary") can be customized to English or Polish.

    Args:
        note_title (str): The title of the note to be included in the document.
        note_summary (str): A brief summary of the note.
        note_content (list): A list of dictionaries representing the note content. Each dictionary should have:
            - 'type' (str): The type of content ('img', 'speaker', or 'text').
            - 'timestamp' (str): The timestamp associated with the content.
            - 'file_path' (str, optional): Path to the image file (for 'img' type only).
            - 'name' (str, optional): The speaker's name (for 'speaker' type only).
            - 'value' (str, optional): The text content (for 'text' type only).
        docx_file_path (str): The file path where the generated DOCX file will be saved.
        language (str, optional): The language for headings ('pl' for Polish, 'en' for English). Defaults to 'pl'.

    Returns:
        bool: `True` if the DOCX file is created and saved successfully. `False` if an error occurs.

    Raises:
        Exception: If there is an error during the DOCX file creation or saving process.

    Error Handling:
        - Logs file operation errors using `log_operations_on_file`.

    Notes:
        - Requires the `python-docx` library.
        - Timestamps are formatted using the `format_timestamp` function.
        - Images are resized to a width of 4 inches using `Inches` from `python-docx`.

    Example:
        >>> content = [
        ...     {'type': 'text', 'timestamp': 0, 'value': 'Introduction to the topic.'},
        ...     {'type': 'speaker', 'timestamp': 10, 'name': 'John Doe'},
        ...     {'type': 'img', 'timestamp': 15, 'file_path': 'example.jpg'}
        ... ]
        >>> success = create_docx_file("Meeting Notes", "Summary of the meeting", content, "meeting_notes.docx")
        >>> if success:
        ...     print("DOCX file created successfully.")
        ... else:
        ...     print("Failed to create DOCX file.")
    """
    title = "Title"
    summary = "Summary"
    if language == 'pl':
        title = "Tytuł"
        summary = "Podsumowanie"

    try:
        docx_file = Document()
        docx_file.add_heading(f"{title}: {note_title}", level=1)
        docx_file.add_heading(f"{summary}:", level=1)
        docx_file.add_paragraph(note_summary)
        docx_file.add_paragraph()

        for content in note_content:
            if content['type'] == 'img':
                docx_file.add_paragraph(format_timestamp(content['timestamp']))
                docx_file.add_picture(content['file_path'], width=Inches(4))
            elif content['type'] == 'speaker':
                docx_file.add_paragraph(f"{format_timestamp(content['timestamp'])} {content['name']}:")
            elif content['type'] == 'text':
                docx_file.add_paragraph(f"{format_timestamp(content['timestamp'])} {content['value']}")

        docx_file.save(docx_file_path)
        return True
    except Exception as e:
        log_operations_on_file(f"[ERROR] create_docx_file(): {repr(e)}")
    return False


def create_txt_file(note_title: str, note_summary: str, note_content: list, txt_file_path: str, language: str ='pl') -> bool:
    """
    Creates a TXT file with the specified note details.

    This function generates a plain text file containing a note's title, summary, and detailed content.
    The content can include speaker labels and text entries, formatted with timestamps. The language of headings
    (e.g., "Title" and "Summary") can be customized to English or Polish.

    Args:
        note_title (str): The title of the note to be included in the document.
        note_summary (str): A brief summary of the note.
        note_content (list): A list of dictionaries representing the note content. Each dictionary should have:
            - 'type' (str): The type of content ('speaker' or 'text').
            - 'timestamp' (str): The timestamp associated with the content.
            - 'name' (str, optional): The speaker's name (for 'speaker' type only).
            - 'value' (str, optional): The text content (for 'text' type only).
        txt_file_path (str): The file path where the generated TXT file will be saved.
        language (str, optional): The language for headings ('pl' for Polish, 'en' for English). Defaults to 'pl'.

    Returns:
        bool: `True` if the TXT file is created and saved successfully. `False` if an error occurs.

    Raises:
        Exception: If there is an error during the TXT file creation or saving process.

    Error Handling:
        - Logs file operation errors using `log_operations_on_file`.

    Notes:
        - Timestamps are formatted using the `format_timestamp` function.

    Example:
        >>> content = [
        ...     {'type': 'text', 'timestamp': 10, 'value': 'Introduction to the topic.'},
        ...     {'type': 'speaker', 'timestamp': 20, 'name': 'John Doe'}
        ... ]
        >>> success = create_txt_file("Meeting Notes", "Summary of the meeting", content, "meeting_notes.txt")
        >>> if success:
        ...     print("TXT file created successfully.")
        ... else:
        ...     print("Failed to create TXT file.")
    """
    title = "Title"
    summary = "Summary"
    if language == 'pl':
        title = "Tytuł"
        summary = "Podsumowanie"

    try:
        with open(txt_file_path, "w", encoding='utf-8') as txt_file:
            txt_file.write(f"{title}: {note_title}\n")
            txt_file.write(f"{summary}:\n")
            txt_file.write(f"{note_summary}\n\n")

            for content in note_content:
                if content['type'] == 'speaker':
                    txt_file.write(f"{format_timestamp(content['timestamp'])} {content['name']}:\n")
                elif content['type'] == 'text':
                    txt_file.write(f"{format_timestamp(content['timestamp'])} {content['value']}\n")
        return True
    except Exception as e:
        log_operations_on_file(f"[ERROR] create_txt_file(): {repr(e)}")
    return False


def create_json_file(note_title: str, note_summary: str, note_content: list, note_datetime: str, video_file_name: str, docx_file_name: str, txt_file_name: str, json_file_path: str, language='pl') -> bool:
    """
    Creates a JSON file containing the specified note details.

    This function generates a JSON file with a note's title, summary, datetime, and content. The content can include
    text, speaker labels, and images, formatted with timestamps. The function also includes file names for related
    video, DOCX, and TXT files. The language of the note (e.g., "Title" and "Summary") can be customized to English or Polish.

    Args:
        note_title (str): The title of the note.
        note_summary (str): A brief summary of the note.
        note_content (list): A list of dictionaries representing the note content. Each dictionary should contain:
            - 'type' (str): The type of content ('img', 'speaker', or 'text').
            - 'timestamp' (str): The timestamp associated with the content.
            - 'file_path' (str, optional): The image file path (for 'img' type only).
            - 'name' (str, optional): The speaker's name (for 'speaker' type only).
            - 'value' (str, optional): The text content (for 'text' type only).
        note_datetime (str): The datetime of the note.
        video_file_name (str): The name of the associated video file.
        docx_file_name (str): The name of the associated DOCX file.
        txt_file_name (str): The name of the associated TXT file.
        json_file_path (str): The path where the generated JSON file will be saved.
        language (str, optional): The language for the note (defaults to 'pl' for Polish).

    Returns:
        bool: `True` if the JSON file is created and saved successfully. `False` if an error occurs.

    Raises:
        Exception: If there is an error during the JSON file creation or saving process.

    Error Handling:
        - Logs file operation errors using `log_operations_on_file`.

    Notes:
        - Timestamps are formatted using the `format_timestamp` function.
        - The JSON file will include the note's ID, derived from the file name (excluding the `.json` extension).

    Example:
        >>> content = [
        ...     {'type': 'text', 'timestamp': 0, 'value': 'Introduction to the topic.'},
        ...     {'type': 'speaker', 'timestamp': 10, 'name': 'John Doe'},
        ...     {'type': 'img', 'timestamp': 10, 'file_path': 'example.jpg'}
        ... ]
        >>> success = create_json_file("Meeting Notes", "Summary of the meeting", content, "2025-01-15", "video.mp4", "meeting_notes.docx", "meeting_notes.txt", "meeting_notes.json")
        >>> if success:
        ...     print("JSON file created successfully.")
        ... else:
        ...     print("Failed to create JSON file.")
    """
    data_content = []
    for content in note_content:
        if content['type'] == 'img':
            data_content.append({"type": "img", "val": f"{os.path.basename(content['file_path'])}", "timestamp_str": f"{format_timestamp(content['timestamp'])}"})
        elif content['type'] == 'speaker':
            data_content.append({"type": "speaker", "val": f"{content['name']}", "timestamp_str": f"{format_timestamp(content['timestamp'])}"})
        elif content['type'] == 'text':
            data_content.append({"type": "txt", "val": f"{content['value']}", "timestamp_str": f"{format_timestamp(content['timestamp'])}"})

    data = {
        "note_id": os.path.basename(json_file_path)[:-5], # -5 -> without .json
        "title": note_title,
        "datetime": note_datetime,
        "summary": note_summary,
        "language": language,
        "video": video_file_name,
        "docx": docx_file_name,
        "txt": txt_file_name,
        "content": data_content
    }

    try:
        with open(json_file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        return True
    except Exception as e:
        log_operations_on_file(f"[ERROR] create_json_file(): {repr(e)}")
    return False